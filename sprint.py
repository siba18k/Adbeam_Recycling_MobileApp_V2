from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# Create a new Document
doc = Document()

# Set up styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title Page
title = doc.add_heading('Adbeam Recycling Mobile App', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.runs[0]
title_run.font.color.rgb = RGBColor(16, 185, 129)  # Green color

subtitle = doc.add_paragraph('Sprint Documentation: Current Progress vs. Planned Features')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_format = subtitle.runs[0]
subtitle_format.font.size = Pt(14)
subtitle_format.font.bold = True

doc.add_paragraph()

# Project Info
info_table = doc.add_table(rows=8, cols=2)
info_table.style = 'Light Grid Accent 1'

info_data = [
    ('Project Name:', 'Adbeam Recycling Mobile App'),
    ('Group Name:', 'Abeam Corporation'),
    ('Group Leader:', 'ST DUBE (223003057)'),
    ('Team Members:', 'L MBOKAZI (223153718), AD MNAMATELI (223029043)'),
    ('', 'NW DLAMINI (224019401), L.S.M MASALESA (223014114)'),
    ('', 'BG SIMANGO (224095653)'),
    ('Document Date:', datetime.now().strftime('%B %d, %Y')),
    ('Sprint Version:', 'Version 1.0')
]

for i, (label, value) in enumerate(info_data):
    row = info_table.rows[i]
    row.cells[0].text = label
    row.cells[1].text = value
    if label:
        row.cells[0].paragraphs[0].runs[0].font.bold = True

doc.add_page_break()

# Executive Summary
doc.add_heading('Executive Summary', 1)
summary = doc.add_paragraph(
    'This document provides a comprehensive comparison between the current implementation status '
    'of the Adbeam Recycling Mobile App and the planned features outlined in the project proposal. '
    'The Adbeam Mobile App is designed to tackle pollution at the college level by incentivizing '
    'students to recycle through a mobile-first rewards platform.'
)

doc.add_paragraph()

# Current Implementation Status
doc.add_heading('1. Current Implementation Status', 1)

doc.add_heading('1.1 Technical Foundation', 2)
tech_para = doc.add_paragraph()
tech_para.add_run('Platform: ').bold = True
tech_para.add_run('React Native with Expo\n')
tech_para.add_run('Status: ').bold = True
tech_para.add_run('✓ Implemented\n').font.color.rgb = RGBColor(16, 185, 129)
tech_para.add_run('Details: ').bold = True
tech_para.add_run('The app is built using React Native and Expo, providing cross-platform compatibility for iOS and Android devices.')

doc.add_heading('1.2 Implemented Features', 2)

# Create a table for implemented features
impl_table = doc.add_table(rows=1, cols=3)
impl_table.style = 'Light Grid Accent 1'

# Header row
header_cells = impl_table.rows[0].cells
header_cells[0].text = 'Feature'
header_cells[1].text = 'Status'
header_cells[2].text = 'Description'

for cell in header_cells:
    cell.paragraphs[0].runs[0].font.bold = True

# Implemented features data
implemented_features = [
    ('Authentication System', '✓ Complete', 'Login and registration screens with Firebase authentication integration'),
    ('Navigation Structure', '✓ Complete', 'Bottom tab navigation with Stack navigator for screen transitions'),
    ('Dashboard Screen', '✓ Complete', 'User dashboard displaying points, scans, and level with mock data'),
    ('Rewards Marketplace', '✓ Complete', 'Browse rewards by category, view details, and check affordability'),
    ('Leaderboard', '✓ Complete', 'Display top recyclers with ranking system'),
    ('Profile Management', '✓ Complete', 'User profile screen with account information'),
    ('Admin Dashboard', '✓ Complete', 'Administrative interface for system management'),
    ('Activity History', '✓ Complete', 'Track and display user recycling activities'),
    ('Firebase Integration', '✓ Complete', 'Backend services for data storage and retrieval'),
    ('Offline Support', '✓ Partial', 'Offline context and storage service implemented'),
    ('User Context', '✓ Complete', 'Global state management for user authentication'),
    ('Reward Categories', '✓ Complete', 'Food, Merchandise, Academic, and Digital categories'),
]

for feature, status, description in implemented_features:
    row = impl_table.add_row()
    row.cells[0].text = feature
    row.cells[1].text = status
    row.cells[2].text = description
    
    # Color the status
    if '✓' in status:
        row.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(16, 185, 129)

doc.add_page_break()

# Not Yet Implemented Features
doc.add_heading('2. Planned Features Not Yet Implemented', 1)

doc.add_heading('2.1 Core Scanning Functionality', 2)

not_impl_table = doc.add_table(rows=1, cols=3)
not_impl_table.style = 'Light Grid Accent 1'

# Header row
header_cells = not_impl_table.rows[0].cells
header_cells[0].text = 'Feature'
header_cells[1].text = 'Priority'
header_cells[2].text = 'Description'

for cell in header_cells:
    cell.paragraphs[0].runs[0].font.bold = True

# Not implemented features
not_implemented = [
    ('Barcode Scanning', 'CRITICAL', 'Camera-based barcode scanning for recyclable items - currently shows placeholder'),
    ('Barcode Validation', 'CRITICAL', 'One-time barcode validation to prevent duplicate submissions'),
    ('Material Recognition', 'HIGH', 'Automatic identification of material type (glass, aluminum, plastic)'),
    ('Points Calculation', 'HIGH', 'Real-time credit calculation based on material type (glass: 10pts, aluminum: 7pts, plastic: 5pts)'),
    ('Camera Optimization', 'HIGH', 'Auto-focus, LED flash integration, and low-light scanning'),
]

for feature, priority, description in not_implemented:
    row = not_impl_table.add_row()
    row.cells[0].text = feature
    row.cells[1].text = priority
    row.cells[2].text = description
    
    # Color priority
    if priority == 'CRITICAL':
        row.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(220, 38, 38)
    elif priority == 'HIGH':
        row.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(245, 158, 11)

doc.add_heading('2.2 Location & Verification Features', 2)

location_table = doc.add_table(rows=1, cols=3)
location_table.style = 'Light Grid Accent 1'

header_cells = location_table.rows[0].cells
header_cells[0].text = 'Feature'
header_cells[1].text = 'Priority'
header_cells[2].text = 'Description'

for cell in header_cells:
    cell.paragraphs[0].runs[0].font.bold = True

location_features = [
    ('GPS Geofencing', 'HIGH', 'Verify recycling activities occur within campus boundaries'),
    ('Location-Based Services', 'MEDIUM', 'Guide students to nearby recycling bins and collection points'),
    ('Campus Mapping', 'MEDIUM', 'Interactive map showing recycling hotspots'),
]

for feature, priority, description in location_features:
    row = location_table.add_row()
    row.cells[0].text = feature
    row.cells[1].text = priority
    row.cells[2].text = description
    
    if priority == 'HIGH':
        row.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(245, 158, 11)
    elif priority == 'MEDIUM':
        row.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(59, 130, 246)

doc.add_heading('2.3 Gamification & Engagement', 2)

gamification_table = doc.add_table(rows=1, cols=3)
gamification_table.style = 'Light Grid Accent 1'

header_cells = gamification_table.rows[0].cells
header_cells[0].text = 'Feature'
header_cells[1].text = 'Priority'
header_cells[2].text = 'Description'

for cell in header_cells:
    cell.paragraphs[0].runs[0].font.bold = True

gamification_features = [
    ('Push Notifications', 'HIGH', 'Daily reminders, achievement alerts, and bonus point notifications'),
    ('Streak Counters', 'MEDIUM', 'Track consecutive recycling days'),
    ('Achievement Badges', 'MEDIUM', 'Unlock badges for recycling milestones'),
    ('Social Sharing', 'MEDIUM', 'Share achievements to social media platforms'),
    ('Challenges & Competitions', 'MEDIUM', 'Campus-wide recycling challenges'),
    ('Progress Visualization', 'LOW', 'Animated progress bars and impact metrics'),
]

for feature, priority, description in gamification_features:
    row = gamification_table.add_row()
    row.cells[0].text = feature
    row.cells[1].text = priority
    row.cells[2].text = description
    
    if priority == 'HIGH':
        row.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(245, 158, 11)
    elif priority == 'MEDIUM':
        row.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(59, 130, 246)
    elif priority == 'LOW':
        row.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(107, 114, 128)

doc.add_page_break()

# Implementation Phases
doc.add_heading('3. Recommended Implementation Phases', 1)

doc.add_heading('Phase 1: Core Functionality (Sprint 2-3)', 2)
phase1 = doc.add_paragraph()
phase1.add_run('Priority: ').bold = True
phase1.add_run('CRITICAL\n')
phase1.add_run('Timeline: ').bold = True
phase1.add_run('2-3 weeks\n\n')
phase1.add_run('Features to Implement:\n').bold = True

phase1_features = [
    'Implement barcode scanning using expo-barcode-scanner',
    'Create barcode validation system to prevent duplicates',
    'Build material type recognition logic',
    'Implement real-time points calculation and awarding',
    'Connect scanner to Firebase for data persistence',
    'Add scanning animations and user feedback',
]

for feature in phase1_features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_heading('Phase 2: Location & Enhanced Features (Sprint 4-5)', 2)
phase2 = doc.add_paragraph()
phase2.add_run('Priority: ').bold = True
phase2.add_run('HIGH\n')
phase2.add_run('Timeline: ').bold = True
phase2.add_run('2-3 weeks\n\n')
phase2.add_run('Features to Implement:\n').bold = True

phase2_features = [
    'Integrate GPS geofencing for campus boundary verification',
    'Implement push notification system',
    'Add location-based recycling bin finder',
    'Enhance offline synchronization capabilities',
    'Implement real-time data updates',
]

for feature in phase2_features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_heading('Phase 3: Gamification & Social Features (Sprint 6-7)', 2)
phase3 = doc.add_paragraph()
phase3.add_run('Priority: ').bold = True
phase3.add_run('MEDIUM\n')
phase3.add_run('Timeline: ').bold = True
phase3.add_run('2-3 weeks\n\n')
phase3.add_run('Features to Implement:\n').bold = True

phase3_features = [
    'Build streak tracking system',
    'Create achievement badge system',
    'Implement social sharing functionality',
    'Add campus challenges and competitions',
    'Enhance leaderboard with filters and categories',
    'Add friend connections and social features',
]

for feature in phase3_features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_heading('Phase 4: Advanced Features (Sprint 8+)', 2)
phase4 = doc.add_paragraph()
phase4.add_run('Priority: ').bold = True
phase4.add_run('LOW\n')
phase4.add_run('Timeline: ').bold = True
phase4.add_run('3-4 weeks\n\n')
phase4.add_run('Features to Implement:\n').bold = True

phase4_features = [
    'Augmented reality for recycling education',
    'Voice command integration',
    'Machine learning for improved item recognition',
    'IoT integration with smart recycling bins',
    'Campus system integrations (dining, events, etc.)',
    'Multi-campus deployment capabilities',
]

for feature in phase4_features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_page_break()

# Technical Debt & Improvements
doc.add_heading('4. Technical Improvements Needed', 1)

improvements = [
    ('Firebase Configuration', 'The Firebase config needs to be properly set up with environment variables for production deployment.'),
    ('Error Handling', 'Enhance error handling throughout the app with user-friendly error messages.'),
    ('Loading States', 'Improve loading states and skeleton screens for better user experience.'),
    ('Data Validation', 'Add comprehensive input validation for all user-submitted data.'),
    ('Security', 'Implement proper security rules for Firebase and secure API endpoints.'),
    ('Testing', 'Add unit tests, integration tests, and end-to-end tests.'),
    ('Performance', 'Optimize image loading, implement lazy loading, and reduce bundle size.'),
    ('Accessibility', 'Add screen reader support, high contrast mode, and voice guidance.'),
]

for title, description in improvements:
    p = doc.add_paragraph()
    p.add_run(f'{title}: ').bold = True
    p.add_run(description)

doc.add_page_break()

# Progress Summary
doc.add_heading('5. Progress Summary', 1)

summary_table = doc.add_table(rows=1, cols=2)
summary_table.style = 'Medium Grid 1 Accent 1'

header_cells = summary_table.rows[0].cells
header_cells[0].text = 'Category'
header_cells[1].text = 'Completion Status'

for cell in header_cells:
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

progress_data = [
    ('App Structure & Navigation', '100% Complete'),
    ('Authentication System', '100% Complete'),
    ('User Interface Screens', '90% Complete'),
    ('Firebase Backend Integration', '80% Complete'),
    ('Core Scanning Functionality', '0% Complete - CRITICAL'),
    ('Location Services', '0% Complete'),
    ('Push Notifications', '0% Complete'),
    ('Gamification Features', '10% Complete'),
    ('Social Features', '0% Complete'),
    ('Advanced Features (AR, ML)', '0% Complete'),
]

for category, status in progress_data:
    row = summary_table.add_row()
    row.cells[0].text = category
    row.cells[1].text = status
    
    # Color code the status
    if '100%' in status:
        row.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(16, 185, 129)
    elif 'CRITICAL' in status or '0%' in status:
        row.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(220, 38, 38)
    else:
        row.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(245, 158, 11)

doc.add_paragraph()

overall = doc.add_paragraph()
overall.add_run('Overall Project Completion: ').bold = True
overall.add_run('~35%').font.color.rgb = RGBColor(245, 158, 11)

doc.add_page_break()

# Recommendations
doc.add_heading('6. Recommendations for Next Sprint', 1)

doc.add_paragraph(
    'Based on the current implementation status, the following recommendations are made for the next sprint:'
)

recommendations = [
    'Prioritize implementing the barcode scanning functionality as it is the core feature of the application',
    'Set up proper Firebase configuration with production credentials',
    'Implement barcode validation to prevent duplicate scans',
    'Add material type recognition and points calculation logic',
    'Create comprehensive testing strategy for scanning functionality',
    'Document the scanning workflow and edge cases',
    'Plan for GPS integration in the following sprint',
    'Begin designing the push notification system architecture',
]

for i, rec in enumerate(recommendations, 1):
    doc.add_paragraph(f'{i}. {rec}')

doc.add_page_break()

# Conclusion
doc.add_heading('7. Conclusion', 1)

conclusion = doc.add_paragraph(
    'The Adbeam Recycling Mobile App has made significant progress in establishing the foundational '
    'architecture and user interface. The authentication system, navigation structure, and basic screens '
    'are fully implemented and functional. However, the core scanning functionality—which is the primary '
    'value proposition of the application—remains unimplemented.\n\n'
    'The next sprint should focus exclusively on implementing the barcode scanning feature, as all other '
    'features depend on this core functionality. Once scanning is operational, the team can proceed with '
    'location services, push notifications, and gamification features in subsequent sprints.\n\n'
    'With focused effort on the critical path items, the Adbeam Mobile App can achieve a functional MVP '
    'within 2-3 sprints, followed by enhancement sprints to add engagement and social features.'
)

doc.add_paragraph()

# Sign-off section
doc.add_paragraph('_' * 50)
doc.add_paragraph()

signoff = doc.add_paragraph()
signoff.add_run('Prepared by: ').bold = True
signoff.add_run('Abeam Corporation Development Team\n')
signoff.add_run('Date: ').bold = True
signoff.add_run(datetime.now().strftime('%B %d, %Y\n'))
signoff.add_run('Version: ').bold = True
signoff.add_run('1.0')

# Save the document
doc.save('/tmp/Adbeam_Sprint_Documentation.docx')

print("Sprint documentation generated successfully!")
print("File saved as: Adbeam_Sprint_Documentation.docx")
